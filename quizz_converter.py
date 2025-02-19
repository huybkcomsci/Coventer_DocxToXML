import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import xml.sax.saxutils as saxutils
import re
from docx import Document
from docx.shared import RGBColor
import os

class QuizzConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("Quiz Converter to Moodle XML")
        self.root.geometry("1200x600")
        
        self.create_widgets()
        self.input_file_path = ""
        self.questions = []  # To store questions dynamically
    
    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Input Section
        input_frame = ttk.LabelFrame(main_frame, text="Đầu vào")
        input_frame.pack(fill=tk.X, pady=5)
        
        ttk.Button(
            input_frame, 
            text="Chọn file câu hỏi",
            command=self.browse_input_file
        ).pack(side=tk.LEFT, padx=5)
        
        self.lbl_input_file = ttk.Label(input_frame, text="Chưa chọn file")
        self.lbl_input_file.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        # Control Section
        control_frame = ttk.Frame(main_frame)
        control_frame.pack(fill=tk.X, pady=10)
        
        ttk.Button(
            control_frame,
            text="Xem trước XML",
            command=self.preview_xml
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            control_frame,
            text="Chuyển đổi",
            command=self.convert_file,
            style="Primary.TButton"
        ).pack(side=tk.RIGHT, padx=5)
        
        # Edit Question Section
        edit_frame = ttk.LabelFrame(main_frame, text="Danh sách câu hỏi")
        edit_frame.pack(fill=tk.X, pady=5)
        
        self.questions_listbox = ttk.Treeview(edit_frame, columns=("ID", "Câu hỏi"), show="headings", height=16)
        self.questions_listbox.heading("ID", text="ID")
        self.questions_listbox.heading("Câu hỏi", text="Câu hỏi")
        self.questions_listbox.column("ID", width=50) 
        self.questions_listbox.column("Câu hỏi", width=400) 
        self.questions_listbox.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Add New Question Button
        ttk.Button(
            edit_frame,
            text="Thêm câu hỏi mới",
            command=self.add_new_question
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            edit_frame,
            text="Chỉnh sửa câu hỏi",
            command=self.edit_question
        ).pack(side=tk.RIGHT, padx=5)
        
        # Log Console
        self.log = scrolledtext.ScrolledText(main_frame, height=10, state=tk.DISABLED)
        self.log.pack(fill=tk.BOTH, expand=True)
        
        # Style Configuration
        style = ttk.Style()
        style.configure("Primary.TButton", foreground="Green", background="#99FF99")
    
    def browse_input_file(self):
        file_types = (
            ('Word files', '*.docx'),
            ('Text files', '*.txt'),
            ('All files', '*.*')
        )
        self.input_file_path = filedialog.askopenfilename(
            title="Chọn file câu hỏi",
            filetypes=file_types
        )
        if self.input_file_path:
            self.lbl_input_file.config(text=self.input_file_path)
            self.log_message(f"Đã chọn file: {self.input_file_path}")
            self.questions = self.parse_file()  # Parse file after selecting it
            self.populate_questions_list()  # Populate the list of questions
    
    def log_message(self, message):
        self.log.config(state=tk.NORMAL)
        self.log.insert(tk.END, message + "\n")
        self.log.see(tk.END)
        self.log.config(state=tk.DISABLED)
    
    def preview_xml(self):
        if not self.input_file_path:
            messagebox.showerror("Lỗi", "Vui lòng chọn file đầu vào trước!")
            return
        
        try:
            xml_content = self.generate_xml_content(self.questions)
            
            preview_window = tk.Toplevel(self.root)
            preview_window.title("Xem trước XML")
            
            text_area = scrolledtext.ScrolledText(
                preview_window, 
                wrap=tk.WORD,
                width=80,
                height=25
            )
            text_area.pack(fill=tk.BOTH, expand=True)
            text_area.insert(tk.INSERT, xml_content)
            text_area.config(state=tk.DISABLED)
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể tạo bản xem trước:\n{str(e)}")
    
    def parse_docx(self, file_path):
        doc = Document(file_path)
        questions = []
        current_question = None
        answer_line = None
        
        all_text = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
        answer_match = re.search(r'(\d+[A-Z]+)+$', all_text)
        if answer_match:
            answer_line = answer_match.group()
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            if re.fullmatch(r'Câu hỏi \d+:.+', text, re.IGNORECASE):
                if current_question:
                    questions.append(current_question)
                
                match = re.match(r'Câu hỏi (\d+):\s*(.*)', text, re.IGNORECASE)
                current_question = {
                    'id': int(match.group(1)),
                    'text': match.group(2),
                    'options': [],
                    'correct': [],
                    'explicit_answer': False
                }
                continue
            
            if current_question:
                option_match = re.fullmatch(r'([A-Z])\.\s*.+', text, re.IGNORECASE)
                if option_match:
                    letter = option_match.group(1).upper()
                    option_text = re.sub(r'^[A-Z]\.\s*', '', text)
                    
                    is_correct = any(self.check_formatting(run) for run in para.runs)
                    if is_correct:
                        current_question['correct'].append(letter)
                    
                    current_question['options'].append((letter, option_text))
                    continue
                
                answer_match = re.match(r'Đáp án đúng:\s*([A-Z,\s]+)', text, re.IGNORECASE)
                if answer_match:
                    answers = re.findall(r'[A-Z]', answer_match.group(1).upper())
                    current_question['correct'] = answers
                    current_question['explicit_answer'] = True
        
        if current_question:
            questions.append(current_question)
        
        if answer_line:
            answer_dict = {}
            matches = re.finditer(r'(\d+)([A-Z]+)', answer_line)
            for match in matches:
                qid = int(match.group(1))
                answers = list(match.group(2).upper())
                answer_dict[qid] = answers
            
            for question in questions:
                if not question['explicit_answer'] and not question['correct']:
                    question['correct'] = answer_dict.get(question['id'], [])
        
        return questions
    
    def parse_file(self):
        if self.input_file_path.endswith('.docx'):
            return self.parse_docx(self.input_file_path)
        else:
            with open(self.input_file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return self.parse_text(content)
    
    def populate_questions_list(self):
        """Populate the list of questions into the Treeview."""
        for row in self.questions_listbox.get_children():
            self.questions_listbox.delete(row)
        
        for question in self.questions:
            self.questions_listbox.insert("", "end", values=(question["id"], question["text"]))
    
    def add_new_question(self):
        add_window = tk.Toplevel(self.root)
        add_window.title("Thêm câu hỏi mới")
        
        # Question Text
        ttk.Label(add_window, text="Câu hỏi:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        question_entry = ttk.Entry(add_window, width=100)
        question_entry.grid(row=0, column=1, padx=5, pady=5)
        
        # Answer Options
        options_frame = ttk.LabelFrame(add_window, text="Lựa chọn câu trả lời")
        options_frame.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky=tk.W+tk.E)
        
        answer_entries = []
        for i in range(4):
            letter = chr(65 + i)
            ttk.Label(options_frame, text=f"{letter}.").grid(row=i, column=0, sticky=tk.W, padx=5)
            answer_entries.append(ttk.Entry(options_frame, width=50))
            answer_entries[i].grid(row=i, column=1, padx=5, pady=5)
        
        # Correct answer(s)
        ttk.Label(add_window, text="Đáp án đúng (vd: A, C):").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        correct_answer_entry = ttk.Entry(add_window, width=50)
        correct_answer_entry.grid(row=2, column=1, padx=5, pady=5)
        
        # Save Button
        def save_new_question():
            question_text = question_entry.get().strip()
            answers = [entry.get().strip() for entry in answer_entries]
            correct_answers = [x.strip().upper() for x in correct_answer_entry.get().split(',')]
            
            if question_text and all(answers) and correct_answers:
                new_question = {
                    'id': len(self.questions) + 1,
                    'text': question_text,
                    'options': [(chr(65 + i), ans) for i, ans in enumerate(answers)],
                    'correct': correct_answers
                }
                self.questions.append(new_question)
                add_window.destroy()
                messagebox.showinfo("Thành công", "Đã thêm câu hỏi mới!")
                self.populate_questions_list()  # Update the list with the new question
            else:
                messagebox.showerror("Lỗi", "Vui lòng điền đầy đủ thông tin.")
        
        ttk.Button(add_window, text="Lưu câu hỏi", command=save_new_question).grid(row=3, column=0, columnspan=2, pady=10)
    
    def edit_question(self):
        selected_item = self.questions_listbox.selection()
        
        if not selected_item:
            messagebox.showerror("Lỗi", "Vui lòng chọn câu hỏi cần chỉnh sửa!")
            return
        
        selected_id = self.questions_listbox.item(selected_item[0], "values")[0]
        question_to_edit = next(q for q in self.questions if q["id"] == int(selected_id))
        
        edit_window = tk.Toplevel(self.root)
        edit_window.title(f"Chỉnh sửa câu hỏi {selected_id}")
        
        # Question Text
        ttk.Label(edit_window, text="Câu hỏi:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        question_entry = ttk.Entry(edit_window, width=100)
        question_entry.insert(0, question_to_edit['text'])
        question_entry.grid(row=0, column=1, padx=5, pady=5)
        
        # Answer Options
        options_frame = ttk.LabelFrame(edit_window, text="Lựa chọn câu trả lời")
        options_frame.grid(row=1, column=0, columnspan=2, padx=5, pady=5, sticky=tk.W+tk.E)
        
        answer_entries = []
        for i, (letter, option) in enumerate(question_to_edit['options']):
            ttk.Label(options_frame, text=f"{letter}.").grid(row=i, column=0, sticky=tk.W, padx=5)
            answer_entries.append(ttk.Entry(options_frame, width=50))
            answer_entries[i].insert(0, option)
            answer_entries[i].grid(row=i, column=1, padx=5, pady=5)
        
        # Correct answer(s)
        ttk.Label(edit_window, text="Đáp án đúng (vd: A, C):").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        correct_answer_entry = ttk.Entry(edit_window, width=50)
        correct_answer_entry.insert(0, ', '.join(question_to_edit['correct']))
        correct_answer_entry.grid(row=2, column=1, padx=5, pady=5)
        
        # Save Button
        def save_edited_question():
            question_text = question_entry.get().strip()
            answers = [entry.get().strip() for entry in answer_entries]
            correct_answers = [x.strip().upper() for x in correct_answer_entry.get().split(',')]
            
            if question_text and all(answers) and correct_answers:
                question_to_edit['text'] = question_text
                question_to_edit['options'] = [(chr(65 + i), ans) for i, ans in enumerate(answers)]
                question_to_edit['correct'] = correct_answers
                edit_window.destroy()
                messagebox.showinfo("Thành công", "Đã cập nhật câu hỏi!")
                self.populate_questions_list()  # Update the list with the edited question
            else:
                messagebox.showerror("Lỗi", "Vui lòng điền đầy đủ thông tin.")
        
        ttk.Button(edit_window, text="Lưu chỉnh sửa", command=save_edited_question).grid(row=3, column=0, columnspan=2, pady=10)

    def check_formatting(self, run):
        if run.font.color.rgb == RGBColor(255, 0, 0):
            return True
        if run.font.bold or run.font.italic or run.font.underline:
            return True
        return False

    def generate_xml_content(self, questions):
        xml_content = '<?xml version="1.0" encoding="UTF-8"?>\n<quiz>\n'
        
        for question in questions:
            xml_content += '    <question type="multichoice">\n'
            xml_content += '        <name>\n'
            xml_content += f'            <text>{saxutils.escape(question["text"])}</text>\n'
            xml_content += '        </name>\n'
            xml_content += '        <questiontext format="html">\n'
            xml_content += f'            <text>{saxutils.escape(question["text"])}</text>\n'
            xml_content += '        </questiontext>\n'
            
            sorted_options = sorted(question['options'], key=lambda x: x[0])
            num_correct = len(question['correct'])
            
            for letter, text in sorted_options:
                if num_correct > 0:
                    fraction = 100.0 / num_correct if letter in question['correct'] else 0.0
                else:
                    fraction = 0.0
                
                fraction_str = f"{fraction:.2f}".rstrip('0').rstrip('.') if fraction % 1 else f"{int(fraction)}"
                xml_content += f'        <answer fraction="{fraction_str}">\n'
                xml_content += f'            <text>{saxutils.escape(text)}</text>\n'
                xml_content += '        </answer>\n'
            
            xml_content += '    </question>\n'
        
        xml_content += '</quiz>'
        return xml_content
    
    def convert_file(self):
        if not self.input_file_path:
            messagebox.showerror("Lỗi", "Vui lòng chọn file đầu vào!")
            return
        
        try:
            # Use the questions that have been added or edited
            xml_content = self.generate_xml_content(self.questions)
            
            # Tạo thư mục OUTPUT
            output_dir = os.path.join(os.path.dirname(self.input_file_path), "OUTPUT")
            os.makedirs(output_dir, exist_ok=True)
            
            # Tạo tên file đầu ra
            base_name = os.path.splitext(os.path.basename(self.input_file_path))[0]
            output_path = os.path.join(output_dir, f"{base_name}_quiz.xml")
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(xml_content)
            
            messagebox.showinfo("Thành công", f"Đã tạo file XML thành công!\n{output_path}")
            self.log_message(f"Đã xuất file thành công: {output_path}")
            os.startfile(output_dir)
        
        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi khi xử lý file:\n{str(e)}")
            self.log_message(f"Lỗi: {str(e)}")

if __name__ == "__main__":
    root = tk.Tk()
    app = QuizzConverter(root)
    root.mainloop()
